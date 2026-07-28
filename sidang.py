import io
import streamlit as st
import pandas as pd
from datetime import date
from docx import Document

# --- FUNKSI PENGGANTI PLACEHOLDER WORD ---
def generate_docx(template_path, replacements):
    doc = Document(template_path)
    
    # Ganti placeholder di paragraf biasa
    for p in doc.paragraphs:
        for key, val in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, str(val))
                
    # Ganti placeholder di dalam tabel
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, val in replacements.items():
                        if key in p.text:
                            p.text = p.text.replace(key, str(val))
                            
    # Simpan ke memory buffer (tanpa perlu buat file temporary di server)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- LAYOUT STREAMLIT ---
st.set_page_config(page_title="Form Nilai Sidang UNISBA", page_icon="🎓", layout="centered")

st.markdown("""
<div style="text-align: center;">
    <h3>PROGRAM STUDI EKONOMI PEMBANGUNAN</h3>
    <h4>FAKULTAS EKONOMI DAN BISNIS UNISBA</h4>
    <hr style="border: 1px solid #1E3A8A; margin-bottom: 20px;">
    <h3 style="color: #1E3A8A;">INPUT & CETAK FORMILIR NILAI SIDANG</h3>
</div>
""", unsafe_allow_html=True)

# Form Identitas
st.subheader("📋 Identitas Ujian")
col1, col2 = st.columns(2)

with col1:
    nama = st.text_input("Nama Mahasiswa", value="Alinda Mutiara Salwa")
    npm = st.text_input("NPM / NIRM", value="10090222040")

with col2:
    nama_penguji = st.text_input("Nama Penguji", value="Yuhka Sundaya, SE., M.Si.")
    tgl_sidang = st.date_input("Tanggal Ujian", value=date(2026, 7, 29))

judul_skripsi = st.text_area(
    "Judul Skripsi", 
    value="ANALISIS FAKTOR-FAKTOR EKONOMI PERILAKU YANG MEMPENGARUHI KETERLIBATAN SHARING KONTEN PRODUK DI TIKTOK",
    height=80
)

st.divider()

# Penilaian
st.subheader("📊 Penilaian Ujian")

col_n1, col_n2, col_n3 = st.columns([3, 2, 2])

with col_n1:
    st.write("**Kriteria**")
    st.write("1. Penyajian (Presentasi)")
    st.write("2. Metodologi")
    st.write("3. Materi")

with col_n2:
    st.write("**Bobot**")
    st.write("20%")
    st.write("30%")
    st.write("50%")

with col_n3:
    st.write("**Nilai (0-100)**")
    np_val = st.number_input("Penyajian", min_value=0.0, max_value=100.0, value=80.0, step=1.0, label_visibility="collapsed")
    nm_val = st.number_input("Metodologi", min_value=0.0, max_value=100.0, value=82.0, step=1.0, label_visibility="collapsed")
    nk_val = st.number_input("Materi", min_value=0.0, max_value=100.0, value=85.0, step=1.0, label_visibility="collapsed")

# Hitung
bp_val = np_val * 0.20
bm_val = nm_val * 0.30
bk_val = nk_val * 0.50
total = bp_val + bm_val + bk_val
status = "lulus" if total >= 60 else "tidak lulus"

st.divider()

# Catatan Revisi
st.subheader("📝 Catatan Revisi")
revisi = st.text_area("Revisi, Perbaikan, dan Catatan Skripsi", height=120)
tgl_revisi = st.text_input("Tanggal Selesai Revisi", value=".......................2026")

# Kamus nilai pengganti
replacements = {
    "{{NAMA}}": nama,
    "{{NPM}}": npm,
    "{{JUDUL}}": judul_skripsi,
    "{{PENGUJI}}": nama_penguji,
    "{{TANGGAL}}": tgl_sidang.strftime("%d %B %Y"),
    "{{NP}}": f"{np_val:.1f}",
    "{{BP}}": f"{bp_val:.2f}",
    "{{NM}}": f"{nm_val:.1f}",
    "{{BM}}": f"{bm_val:.2f}",
    "{{NK}}": f"{nk_val:.1f}",
    "{{BK}}": f"{bk_val:.2f}",
    "{{TOTAL}}": f"{total:.2f}",
    "{{STATUS}}": status.upper(),
    "{{REVISI}}": revisi if revisi else "-",
    "{{TGL_REVISI}}": tgl_revisi
}

st.divider()

# --- TOMBOL GENERATE & DOWNLOAD ---
try:
    docx_file = generate_docx("template_form.docx", replacements)
    
    st.download_button(
        label="📄 Download Form Nilai Sidang (DOCX)",
        data=docx_file,
        file_name=f"Form Nilai Sidang - {nama} - {npm}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary",
        use_container_width=True
    )
except FileNotFoundError:
    st.error("⚠️ File `template_form.docx` belum ditemukan di folder aplikasi. Mohon siapkan file template terlebih dahulu.")
